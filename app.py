import os
from io import BytesIO
from datetime import datetime
from flask import Flask, request, send_file, render_template_string, jsonify

# --- PDF (ReportLab) ---
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# --- DOCX ---
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET", "change-me")

from flask import Flask, render_template, render_template_string

app = Flask(__name__)

# Landing page
@app.route('/')
def landing():
    return render_template('landing.html')

@app.route('/resume', methods=['GET', 'POST'])
def builder():
    return render_template('resume.html')

if __name__ == '__main__':
    app.run(debug=True)

# -------------------------
# HTML (single file: HTML+CSS+JS in one string)
# We wrap the entire template in {% raw %} … {% endraw %} so Jinja doesn’t try to interpret {{ }} inside JS/CSS.
# -------------------------
# -------------------------
# Utility: safe text for PDF/DOCX
# -------------------------
def _safe(s):
    return (s or "").strip()

def _join_nonempty(*items, sep=" — "):
    parts = [p for p in items if _safe(p)]
    return sep.join(parts)

# -------------------------
# PDF generator (classic ATS look)
# -------------------------
def generate_pdf_bytes(data: dict) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54
    )
    styles = getSampleStyleSheet()
    base = styles["Normal"]
    base.fontName = "Helvetica"
    base.fontSize = 10.5
    base.leading = 13

    h_name = ParagraphStyle('h_name', parent=base, fontSize=18, leading=22, spaceAfter=4, alignment=TA_CENTER)
    h_job  = ParagraphStyle('h_job', parent=base, fontSize=12, leading=15, textColor="#1f3c88", spaceAfter=6, alignment=TA_CENTER)
    h_sec  = ParagraphStyle('h_sec', parent=base, fontSize=11, leading=14, spaceBefore=10, spaceAfter=4)
    small  = ParagraphStyle('small', parent=base, fontSize=9.5, leading=12)

    story = []

    name = _safe(data.get("name")) or "Your Name"
    job  = _safe(data.get("job_title")) or "Target Job Title"
    story.append(Paragraph(f"<b>{name}</b>", h_name))
    story.append(Paragraph(job, h_job))

    # Contact
    contact_items = []
    if _safe(data.get("phone")): contact_items.append(_safe(data["phone"]))
    if _safe(data.get("email")): contact_items.append(_safe(data["email"]))
    if _safe(data.get("portfolio_link")): contact_items.append(_safe(data["portfolio_link"]))
    if contact_items:
        story.append(Paragraph(" | ".join(contact_items), small))

    # Address
    addr = [data.get("address1"), data.get("address2"), data.get("place"), data.get("country"), data.get("postalCode")]
    addr = [a for a in addr if _safe(a)]
    if addr:
        story.append(Paragraph("<b>Address</b>", h_sec))
        story.append(Paragraph("<br/>".join(map(_safe, addr)), base))

    # Summary
    if _safe(data.get("summary")):
        story.append(Paragraph("<b>Summary</b>", h_sec))
        story.append(Paragraph(_safe(data["summary"]).replace("\n","<br/>"), base))

    # Experience
    exp = data.get("experience") or []
    if exp:
        story.append(Paragraph("<b>Experience</b>", h_sec))
        for e in exp:
            line = _join_nonempty(_safe(e.get("role")), _safe(e.get("company")))
            story.append(Paragraph(f"<b>{line}</b>", base))
            dates = _join_nonempty(_safe(e.get("start")), _safe(e.get("end")))
            if dates:
                story.append(Paragraph(dates, small))
            ach = e.get("achievements") or []
            for a in ach:
                story.append(Paragraph(f"• { _safe(a) }", base))
            story.append(Spacer(1, 4))

    # Education
    edu = data.get("education") or []
    if edu:
        story.append(Paragraph("<b>Education</b>", h_sec))
        for ed in edu:
            line = _join_nonempty(_safe(ed.get("degree")), _safe(ed.get("school")))
            story.append(Paragraph(f"<b>{line}</b>", base))
            dates = _join_nonempty(_safe(ed.get("start")), _safe(ed.get("end")))
            if dates:
                story.append(Paragraph(dates, small))
            story.append(Spacer(1, 2))

    # Certifications
    certs = data.get("certifications") or []
    if certs:
        story.append(Paragraph("<b>Certifications</b>", h_sec))
        for c in certs:
            line = _join_nonempty(_safe(c.get("name")), _safe(c.get("issuer")))
            story.append(Paragraph(f"{line}", base))
            if _safe(c.get("date")):
                story.append(Paragraph(_safe(c["date"]), small))

    # Skills
    skills = data.get("skills") or []
    if skills:
        story.append(Paragraph("<b>Skills</b>", h_sec))
        story.append(Paragraph(", ".join(map(_safe, skills)), base))

    # Languages
    langs = data.get("languages") or []
    if langs:
        story.append(Paragraph("<b>Languages</b>", h_sec))
        for l in langs:
            story.append(Paragraph(_join_nonempty(_safe(l.get("language")), _safe(l.get("proficiency"))), base))

    # Extras
    extras = data.get("extras") or []
    if extras:
        story.append(Paragraph("<b>Extras</b>", h_sec))
        for x in extras:
            line = _safe(x.get("desc"))
            date = _safe(x.get("date"))
            story.append(Paragraph(_join_nonempty(line, date), base))

    doc.build(story)
    buf.seek(0)
    return buf

# ---------------------------
# DOCX generator
# ---------------------------
def generate_docx_bytes(data: dict) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    name = _safe(data.get("name")) or "Your Name"
    job  = _safe(data.get("job_title")) or "Target Job Title"
    doc.add_heading(name, level=1)
    p = doc.add_paragraph(job)
    p.runs[0].bold = True

    # Contact
    contact = []
    if _safe(data.get("phone")): contact.append(_safe(data["phone"]))
    if _safe(data.get("email")): contact.append(_safe(data["email"]))
    if _safe(data.get("portfolio_link")): contact.append(_safe(data["portfolio_link"]))
    if contact:
        doc.add_paragraph(" | ".join(contact))

    # Address
    addr = [data.get("address1"), data.get("address2"), data.get("place"), data.get("country"), data.get("postalCode")]
    addr = [a for a in addr if _safe(a)]
    if addr:
        doc.add_heading("Address", level=2)
        for a in addr:
            doc.add_paragraph(_safe(a))

    # Summary
    if _safe(data.get("summary")):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(_safe(data["summary"]))

    # Experience
    exp = data.get("experience") or []
    if exp:
        doc.add_heading("Experience", level=2)
        for e in exp:
            line = _join_nonempty(_safe(e.get("role")), _safe(e.get("company")))
            doc.add_paragraph(line, style='List Bullet')
            dates = _join_nonempty(_safe(e.get("start")), _safe(e.get("end")))
            if dates:
                doc.add_paragraph(dates)
            for a in (e.get("achievements") or []):
                doc.add_paragraph(_safe(a), style='List Bullet')

    # Education
    edu = data.get("education") or []
    if edu:
        doc.add_heading("Education", level=2)
        for ed in edu:
            line = _join_nonempty(_safe(ed.get("degree")), _safe(ed.get("school")))
            doc.add_paragraph(line, style='List Bullet')
            dates = _join_nonempty(_safe(ed.get("start")), _safe(ed.get("end")))
            if dates:
                doc.add_paragraph(dates)

    # Certifications
    certs = data.get("certifications") or []
    if certs:
        doc.add_heading("Certifications", level=2)
        for c in certs:
            line = _join_nonempty(_safe(c.get("name")), _safe(c.get("issuer")))
            doc.add_paragraph(line, style='List Bullet')
            if _safe(c.get("date")):
                doc.add_paragraph(_safe(c["date"]))

    # Skills
    skills = data.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(map(_safe, skills)))

    # Languages
    langs = data.get("languages") or []
    if langs:
        doc.add_heading("Languages", level=2)
        for l in langs:
            doc.add_paragraph(_join_nonempty(_safe(l.get("language")), _safe(l.get("proficiency"))), style='List Bullet')

    # Extras
    extras = data.get("extras") or []
    if extras:
        doc.add_heading("Extras", level=2)
        for x in extras:
            doc.add_paragraph(_join_nonempty(_safe(x.get("desc")), _safe(x.get("date"))), style='List Bullet')

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# -------------------------
# Routes
# -------------------------
@app.route("/", methods=["GET"])
def index():
    return render_template_string(index)

@app.route("/download-pdf", methods=["POST"])
def download_pdf():
    data = request.get_json(silent=True) or {}
    # Portfolio URL validation (https + LinkedIn/Indeed)
    port = (data.get("portfolio_link") or "").strip()
    if port and not (port.startswith("https://") and ("linkedin.com" in port or "indeed.com" in port)):
        return jsonify({"error": "Portfolio must be a secure LinkedIn or Indeed URL (https://)"}), 400

    pdf_io = generate_pdf_bytes(data)
    filename = f"{(data.get('name') or 'resume').replace(' ','_')}_Resume.pdf"
    return send_file(pdf_io, mimetype="application/pdf", as_attachment=True, download_name=filename)

@app.route("/download-docx", methods=["POST"])
def download_docx():
    data = request.get_json(silent=True) or {}
    # Portfolio URL validation (https + LinkedIn/Indeed)
    port = (data.get("portfolio_link") or "").strip()
    if port and not (port.startswith("https://") and ("linkedin.com" in port or "indeed.com" in port)):
        return jsonify({"error": "Portfolio must be a secure LinkedIn or Indeed URL (https://)"}), 400

    docx_io = generate_docx_bytes(data)
    filename = f"{(data.get('name') or 'resume').replace(' ','_')}_Resume.docx"
    return send_file(docx_io,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=filename)

if __name__ == "__main__":
    # Install deps (once): pip install flask reportlab python-docx
    app.run(debug=True, host="0.0.0.0", port=int(os.getenv("PORT", 5000)))
